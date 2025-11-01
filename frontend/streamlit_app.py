import streamlit as st
import requests
import io

# Optional dependencies for file parsing
try:
    from docx import Document
except ImportError:
    st.warning("python-docx not installed. DOCX parsing won't work.")
try:
    import PyPDF2
except ImportError:
    st.warning("PyPDF2 not installed. PDF parsing won't work.")

# ---------------------------
# Config
# ---------------------------
API_BASE = st.secrets.get("API_BASE", "http://127.0.0.1:8000")
st.set_page_config(
    page_title="Resume Screening & Interview QGen",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------
# Helper function to extract text
# ---------------------------
def extract_text(file):
    text = ""
    try:
        if file.type == "text/plain":
            try:
                text = file.getvalue().decode("utf-8")
            except UnicodeDecodeError:
                text = file.getvalue().decode("latin1")
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            file.seek(0)
            doc = Document(io.BytesIO(file.getvalue()))
            text = "\n".join([p.text for p in doc.paragraphs])
        elif file.type == "application/pdf":
            file.seek(0)
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.getvalue()))
            text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
    except Exception as e:
        st.warning(f"Could not read file {file.name}: {e}")
    return text

# ---------------------------
# Sidebar: Job Description Input
# ---------------------------
st.sidebar.header("Job Description (JD)")
jd_file = st.sidebar.file_uploader("Upload JD (txt, pdf, docx)", type=["txt", "pdf", "docx"])
jd_text_input = st.sidebar.text_area("Or paste JD text here", height=200)

jd_text = jd_text_input
if jd_file:
    jd_file_text = extract_text(jd_file)
    if jd_file_text:
        jd_text = jd_file_text
        st.sidebar.success(f"JD Loaded: {jd_file.name}")
    else:
        st.sidebar.warning("Could not extract JD. Please paste text manually.")

# ---------------------------
# Main App Layout
# ---------------------------
st.title("🚀 Resume Screening & Interview Question Generator")

# ---------------------------
# Step 1: Upload Resume
# ---------------------------
st.header("Step 1: Upload Candidate Resume")
resume_file = st.file_uploader("Upload Resume (txt, pdf, docx)", type=["txt", "pdf", "docx"])
resume_text = ""
if resume_file:
    resume_text = extract_text(resume_file)
    files = {"file": (resume_file.name, resume_file.getvalue())}
    with st.spinner("Uploading resume to backend..."):
        resp = requests.post(f"{API_BASE}/resume/upload", files=files)
    if resp.status_code == 200:
        data = resp.json()
        st.success(f"Resume Uploaded. Resume ID: {data['resume_id']}")
        st.session_state["resume_id"] = data["resume_id"]
    else:
        st.error(f"Upload failed: {resp.text}")

# ---------------------------
# Step 2: Show Resume & JD side by side
# ---------------------------
if resume_text or jd_text:
    st.header("Step 2: Preview & Compare")
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Resume Preview")
        st.text_area("Resume", value=resume_text, height=300)
    with col2:
        st.subheader("Job Description Preview")
        st.text_area("JD", value=jd_text, height=300)

# ---------------------------
# Step 3: Generate Questions
# ---------------------------
st.header("Step 3: Generate Interview Questions")

if st.button("Generate Questions"):
    resume_id = st.session_state.get("resume_id")
    if not resume_id:
        st.error("Please upload a candidate resume first.")
    elif not jd_text:
        st.error("Please provide a job description.")
    else:
        payload = {"resume_id": resume_id, "query": jd_text}
        with st.spinner("Generating interview questions..."):
            resp = requests.post(f"{API_BASE}/qa/generate", json=payload)

        if resp.status_code == 200:
            out = resp.json()
            st.success("✅ Question generation complete!")

            # Try to extract relevant data
            relevance = None
            tech_qs, beh_qs, sources = [], [], []

            # 1️⃣ If backend returns a list of dicts
            if isinstance(out, list):
                for item in out:
                    if isinstance(item, dict):
                        if "match_score" in item:
                            relevance = item["match_score"]
                        if "q_samples" in item:
                            # Split by pipe or sentence
                            text = item["q_samples"]
                            parts = [p.strip() for p in text.split("|") if p.strip()]
                            for q in parts:
                                if q.lower().startswith("technical"):
                                    tech_qs.append(q)
                                elif q.lower().startswith("behavioral"):
                                    beh_qs.append(q)
                                else:
                                    tech_qs.append(q)
            # 2️⃣ If backend returns a dict with structured keys
            elif isinstance(out, dict):
                relevance = out.get("match_score") or out.get("relevance_score")
                tech_qs = out.get("technical_questions", [])
                beh_qs = out.get("behavioral_questions", [])
                sources = out.get("sources", [])

            # Display Relevance Score
            if relevance is not None:
                st.metric("Resume–JD Match Score", f"{float(relevance)*100:.1f}%")

            # Display Technical Questions
            if tech_qs:
                st.subheader("🧠 Technical Questions")
                for i, q in enumerate(tech_qs, 1):
                    st.markdown(f"{i}. {q}")
            else:
                st.info("No technical questions found.")

            # Display Behavioral Questions
            if beh_qs:
                st.subheader("💬 Behavioral Questions")
                for i, q in enumerate(beh_qs, 1):
                    st.markdown(f"{i}. {q}")
            else:
                st.info("No behavioral questions found.")

            # Sources / Context (where the LLM drew info from)
            if sources:
                st.subheader("📚 Sources / Context")
                for s in sources:
                    st.markdown(f"- {s}")
            else:
                st.markdown(
                    "**Sources / Context:** Information about which parts of the resume or JD were used for generating these questions."
                )
        else:
            st.error("Generation failed: " + resp.text)


# import streamlit as st
# import requests
# from pathlib import Path

# API_BASE = st.secrets.get("API_BASE", "http://127.0.0.1:8000")

# st.set_page_config(page_title="Resume Screening Agent")

# st.title("Resume Screening & QGen Agent")

# st.markdown("Upload a resume (text) — backend will parse and later generate interview questions.")

# uploaded_file = st.file_uploader("Upload resume (txt)", type=["txt", "pdf", "docx"])
# if uploaded_file:
#     files = {"file": (uploaded_file.name, uploaded_file.getvalue())}
#     resp = requests.post(f"{API_BASE}/resume/upload", files=files)
#     if resp.status_code == 200:
#         data = resp.json()
#         st.success(f"Uploaded. Resume ID: {data['resume_id']}")
#         st.text_area("Preview", value=data["preview"], height=200)
#         st.session_state["resume_id"] = data["resume_id"]
#     else:
#         st.error(f"Upload failed: {resp.text}")

# jd_text = st.text_area("Paste Job Description (for matching & question generation)", height=200)

# if st.button("Generate Interview Questions"):
#     resume_id = st.session_state.get("resume_id")
#     if not resume_id:
#         st.error("Upload a resume first.")
#     else:
#         resp = requests.post(f"{API_BASE}/qa/generate", json={"resume_id": resume_id, "query": jd_text})
#         if resp.status_code == 200:
#             out = resp.json()
#             st.write("Answer:", out["answer"])
#             st.write("Sources:", out["sources"])
#         else:
#             st.error("Generation failed: " + resp.text)


# import streamlit as st
# import requests
# import io
# from pathlib import Path

# # Optional dependencies for file parsing
# try:
#     from docx import Document
# except ImportError:
#     st.warning("python-docx not installed. DOCX parsing won't work.")
# try:
#     import PyPDF2
# except ImportError:
#     st.warning("PyPDF2 not installed. PDF parsing won't work.")

# # ---------------------------
# # Config
# # ---------------------------
# API_BASE = st.secrets.get("API_BASE", "http://127.0.0.1:8000")
# st.set_page_config(
#     page_title="Resume Screening & QGen Agent",
#     layout="wide",
#     initial_sidebar_state="expanded",
# )

# # ---------------------------
# # Sidebar - JD Upload/Input
# # ---------------------------
# st.sidebar.title("Job Description (JD) Upload / Input")
# jd_file = st.sidebar.file_uploader("Upload JD (txt, pdf, docx)", type=["txt", "pdf", "docx"])
# jd_text_input = st.sidebar.text_area("Or paste JD text here", height=200)

# # Function to extract text from files safely
# def extract_text(file):
#     text = ""
#     try:
#         # TXT files
#         if file.type == "text/plain":
#             try:
#                 text = file.getvalue().decode("utf-8")
#             except UnicodeDecodeError:
#                 text = file.getvalue().decode("latin1")  # fallback encoding
#         # DOCX files
#         elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             file.seek(0)
#             doc = Document(io.BytesIO(file.getvalue()))
#             text = "\n".join([p.text for p in doc.paragraphs])
#         # PDF files
#         elif file.type == "application/pdf":
#             file.seek(0)
#             pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.getvalue()))
#             text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
#     except Exception as e:
#         st.warning(f"Could not read file {file.name}: {e}")
#     return text

# # Load JD text
# jd_text = jd_text_input
# if jd_file:
#     jd_text_file = extract_text(jd_file)
#     if jd_text_file:
#         jd_text = jd_text_file
#         st.sidebar.success(f"JD Loaded: {jd_file.name}")
#     else:
#         st.sidebar.warning("Could not extract JD from the uploaded file. Please paste text manually.")

# # ---------------------------
# # Main App Layout
# # ---------------------------
# st.title("🚀 Resume Screening & Interview Question Generator")
# st.markdown(
#     """
# This application allows you to:
# - Upload candidate resumes (PDF, DOCX, TXT)
# - Parse resumes into structured data
# - Compare with job descriptions
# - Generate personalized technical and behavioral interview questions
# """
# )

# # ---------------------------
# # Resume Upload Section
# # ---------------------------
# st.header("Step 1: Upload Candidate Resume")
# resume_file = st.file_uploader("Upload Resume (txt, pdf, docx)", type=["txt", "pdf", "docx"])
# if resume_file:
#     files = {"file": (resume_file.name, resume_file.getvalue())}
#     with st.spinner("Uploading and parsing resume..."):
#         resp = requests.post(f"{API_BASE}/resume/upload", files=files)
#     if resp.status_code == 200:
#         data = resp.json()
#         st.success(f"✅ Resume Uploaded. Resume ID: {data['resume_id']}")
#         st.text_area("Resume Preview", value=data.get("preview", ""), height=200)
#         st.session_state["resume_id"] = data["resume_id"]
#     else:
#         st.error(f"Upload failed: {resp.text}")

# # ---------------------------
# # JD Matching & Question Generation
# # ---------------------------
# st.header("Step 2: Generate Interview Questions")
# if st.button("Generate Questions"):
#     resume_id = st.session_state.get("resume_id")
#     if not resume_id:
#         st.error("Please upload a candidate resume first.")
#     elif not jd_text:
#         st.error("Please provide a job description via upload or text input.")
#     else:
#         payload = {"resume_id": resume_id, "query": jd_text}
#         with st.spinner("Generating interview questions..."):
#             resp = requests.post(f"{API_BASE}/qa/generate", json=payload)
#         if resp.status_code == 200:
#             out = resp.json()
#             # Display structured output
#             st.subheader("Technical Questions")
#             for q in out.get("technical_questions", []):
#                 st.markdown(f"- {q}")
#             st.subheader("Behavioral Questions")
#             for q in out.get("behavioral_questions", []):
#                 st.markdown(f"- {q}")
#             st.subheader("Sources / Context")
#             st.write(out.get("sources", "No sources provided"))
#             st.success("✅ Question generation complete!")
#         else:
#             st.error("Generation failed: " + resp.text)

# # ---------------------------
# # Optional Features: Memory & Tool Calls
# # ---------------------------
# st.header("Step 3: Optional Memory & Tool Calls")
# st.markdown(
#     """
# - This system can maintain context across multiple resumes or JD comparisons.
# - Optional integrations: skill lookup APIs, market trend tools, etc.
# """
# )