import pdfplumber
from docx import Document
import io

def read_text_file(file_path: str) -> str:
    """
    Extract text content from PDF, DOCX, or TXT.
    """
    if file_path.endswith(".pdf"):
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text.strip()

    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs]).strip()

    elif file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()

    else:
        raise ValueError("Unsupported file format. Only .pdf, .docx, .txt allowed.")



def extract_text_from_file(uploaded_file):
    """
    Extract text content from uploaded resume or JD file (.txt, .docx, .pdf).
    """
    filename = uploaded_file.filename.lower()
    content = ""

    try:
        if filename.endswith(".txt"):
            content = uploaded_file.file.read().decode("utf-8", errors="ignore")

        elif filename.endswith(".docx"):
            document = Document(uploaded_file.file)
            content = "\n".join([para.text for para in document.paragraphs])

        elif filename.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(uploaded_file.file.read())) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n"

        else:
            content = "Unsupported file format. Please upload .txt, .docx, or .pdf."

    except Exception as e:
        content = f"Error reading file: {str(e)}"

    return content.strip()